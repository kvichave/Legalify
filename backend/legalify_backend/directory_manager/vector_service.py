import os
import logging
from qdrant_client import QdrantClient
from sentence_transformers import SentenceTransformer
import PyPDF2
import docx
from langchain_qdrant import QdrantVectorStore, FastEmbedSparse
from langchain_core.documents import Document
from langchain_core.embeddings import Embeddings
from langchain_classic.retrievers import EnsembleRetriever

logger = logging.getLogger(__name__)

QDRANT_HOST = os.environ.get("QDRANT_HOST", "localhost")
QDRANT_PORT = int(os.environ.get("QDRANT_PORT", 6333))
EMBEDDING_MODEL = "all-MiniLM-L6-v2"

_client = None
_model = None


class SentenceTransformerEmbeddings(Embeddings):
    def __init__(self, model_name: str):
        self.model = SentenceTransformer(model_name)

    def embed_documents(self, texts: list[str]) -> list[list[float]]:
        return self.model.encode(texts, show_progress_bar=False).tolist()

    def embed_query(self, text: str) -> list[float]:
        return self.model.encode([text], show_progress_bar=False)[0].tolist()

    def encode(self, texts, **kwargs):
        return self.model.encode(texts, **kwargs)


def get_qdrant_client():
    global _client
    if _client is None:
        _client = QdrantClient(host=QDRANT_HOST, port=QDRANT_PORT)
    return _client


def get_embedding_model():
    global _model
    if _model is None:
        _model = SentenceTransformerEmbeddings(EMBEDDING_MODEL)
    return _model


def extract_text_from_file(file_path):
    """Extract text content from PDF or DOCX files."""
    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext == ".pdf":
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() or ""
                return text
        elif ext == ".docx":
            doc = docx.Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        elif ext == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        else:
            logger.warning(f"Unsupported file type: {ext}")
            return ""
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {e}")
        return ""


def split_into_chunks(text, chunk_size=500):
    """Split text into overlapping chunks."""
    words = text.split()
    chunks = []
    for i in range(0, len(words), chunk_size // 4):
        chunk = " ".join(words[i : i + chunk_size])
        if chunk.strip():
            chunks.append(chunk)
    return chunks


def store_document_vector(document_id, file_path, collection_name="legal_documents"):
    """Store document embedding in Qdrant using LangChain hybrid search."""
    try:
        embeddings = get_embedding_model()
        client = get_qdrant_client()

        text = extract_text_from_file(file_path)
        if not text.strip():
            logger.warning(f"No text extracted from {file_path}")
            return None

        chunks = split_into_chunks(text, chunk_size=500)

        documents = []
        for i, chunk in enumerate(chunks):
            doc = Document(
                page_content=chunk,
                metadata={
                    "document_id": str(document_id),
                    "chunk_index": i,
                    "file_path": file_path,
                },
            )
            documents.append(doc)

        vector_store = QdrantVectorStore.from_documents(
            documents=documents,
            embedding=embeddings,
            sparse_embedding=FastEmbedSparse(),
            client=client,
            collection_name=collection_name,
        )

        logger.info(
            f"Stored {len(documents)} chunks for document {document_id} in Qdrant"
        )
        return True

    except Exception as e:
        logger.error(f"Error storing vector in Qdrant: {e}")
        raise


def delete_document_vectors(document_id, collection_name):
    """Delete all vector points for a document from Qdrant."""
    try:
        client = get_qdrant_client()

        offset = None
        while True:
            scroll_result = client.scroll(
                collection_name=collection_name,
                scroll_filter=None,
                limit=100,
                with_payload=True,
                offset=offset,
            )

            points_to_delete = []
            for point in scroll_result[0]:
                if point.payload and point.payload.get("document_id") == str(
                    document_id
                ):
                    points_to_delete.append(point.id)

            if points_to_delete:
                client.delete(
                    collection_name=collection_name,
                    points_selector=points_to_delete,
                )

            if scroll_result[1] is None:
                break
            offset = scroll_result[1]

        logger.info(f"Deleted vectors for document {document_id}")
        return True
    except Exception as e:
        logger.error(f"Error deleting vectors from Qdrant: {e}")
        raise


def search_vectors(collection_name, query_embedding, query_text="", limit=5):
    """Hybrid search using LangChain EnsembleRetriever with dense and sparse retrievers."""
    try:
        embeddings = get_embedding_model()
        client = get_qdrant_client()

        dense_vector_store = QdrantVectorStore(
            client=client,
            collection_name=collection_name,
            embedding=embeddings,
        )

        sparse_vector_store = QdrantVectorStore(
            client=client,
            collection_name=collection_name,
            embedding=FastEmbedSparse(),
        )

        dense_retriever = dense_vector_store.as_retriever(search_kwargs={"k": limit})
        sparse_retriever = sparse_vector_store.as_retriever(search_kwargs={"k": limit})

        ensemble_retriever = EnsembleRetriever(
            retrievers=[dense_retriever, sparse_retriever],
            weights=[0.5, 0.5],
        )

        results = ensemble_retriever.invoke(query_text)

        chunks = []
        for doc in results:
            chunks.append(
                {
                    "text": doc.page_content,
                    "file_path": doc.metadata.get("file_path", ""),
                    "score": 1.0,
                }
            )

        return chunks
    except Exception as e:
        logger.error(f"Error searching vectors in {collection_name}: {e}")
        return []
